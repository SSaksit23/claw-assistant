"""
Document Parser Service.

Parses uploaded files (CSV, Excel, PDF, DOCX, text) and extracts
structured expense data: group code, travel date, size, price, etc.

Uses OpenAI for unstructured text extraction when the file format
is not a simple table (e.g., PDF or DOCX with free-form text).
"""

import os
import re
import json
import logging
from typing import Optional

import pandas as pd
from openai import OpenAI

from config import Config

logger = logging.getLogger(__name__)


EXPENSE_FIELDS = {
    "tour_code": "Tour/group code (e.g., GO1TAO5NTAOQW260304)",
    "program_code": "Program code for the travel program",
    "travel_date": "Travel date or date range (e.g., 0304-0309 or 10-13 Mar)",
    "travel_date_start": "Start date of travel (dd/mm/yyyy)",
    "travel_date_end": "End date of travel (dd/mm/yyyy)",
    "pax": "Number of passengers / group size",
    "unit_price": "Price per person/unit",
    "quantity": "Extra multiplier beyond pax (e.g., number of days, nights)",
    "amount": "Total expense amount (unit_price x pax x quantity)",
    "currency": "Detected currency (CNY, THB, USD, etc.)",
    "supplier_name": "Supplier / company name (Party A / 甲方)",
    "description": "Original description text (Chinese or English as-is)",
    "expense_label": "English name for expense (e.g., Airline Ticket, Tour Fare, Guide Fee)",
    "charge_type": "Type: flight, land_tour, single_supplement, service_fee, guide_tip, visa, accommodation, meal, transport, insurance, entrance_fee, commission, other",
    "calculation_note": "Human-readable calculation string (e.g., '2,380 x 20 pax = 47,600')",
}


def parse_file(file_path: str) -> dict:
    """
    Parse any supported file and extract expense records.

    Returns:
        {
            "status": "success" | "error",
            "file_type": "csv" | "excel" | "pdf" | "docx" | "text",
            "records": [ { field: value, ... }, ... ],
            "raw_text": "original text (for non-tabular files)",
            "field_mapping": { "original_col": "mapped_field", ... },
            "errors": [ "any parsing errors" ]
        }
    """
    if not os.path.exists(file_path):
        return {"status": "error", "errors": [f"File not found: {file_path}"], "records": []}

    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Parsing file: {file_path} (type: {ext})")

    try:
        if ext == ".csv":
            return _parse_csv(file_path)
        elif ext in (".xlsx", ".xls"):
            return _parse_excel(file_path)
        elif ext == ".pdf":
            return _parse_pdf(file_path)
        elif ext == ".docx":
            return _parse_docx(file_path)
        elif ext in (".txt", ".md"):
            return _parse_text(file_path)
        else:
            return {"status": "error", "errors": [f"Unsupported file type: {ext}"], "records": []}

    except Exception as e:
        logger.error(f"File parsing failed: {e}", exc_info=True)
        return {"status": "error", "errors": [str(e)], "records": []}


# ---------------------------------------------------------------------------
# CSV parsing
# ---------------------------------------------------------------------------
def _parse_csv(file_path: str) -> dict:
    """Parse CSV file with Thai or English column names."""
    df = pd.read_csv(file_path, encoding="utf-8-sig")
    return _dataframe_to_records(df, "csv")


# ---------------------------------------------------------------------------
# Excel parsing
# ---------------------------------------------------------------------------
def _parse_excel(file_path: str) -> dict:
    """Parse Excel file."""
    df = pd.read_excel(file_path)
    return _dataframe_to_records(df, "excel")


# ---------------------------------------------------------------------------
# DataFrame -> Records (shared by CSV / Excel)
# ---------------------------------------------------------------------------
COLUMN_MAP = {
    # Thai -> English field mapping
    "รหัสทัวร์": "tour_code",
    "รหัสโปรแกรม": "program_code",
    "รหัสกรุ๊ป": "tour_code",
    "โปรแกรมทัวร์": "program_code",
    "วันที่เดินทาง": "travel_date",
    "วันที่": "travel_date",
    "จำนวนลูกค้า หัก หนท.": "pax",
    "จำนวนลูกค้า": "pax",
    "จำนวนคน": "pax",
    "ยอดเบิก": "amount",
    "จำนวนเงิน": "amount",
    "ราคา": "amount",
    "คำอธิบาย": "description",
    "ประเภท": "charge_type",
    "สกุลเงิน": "currency",
    "เรท": "exchange_rate",
    "หมายเหตุ": "remark",
    # Chinese column names
    "编号": "_seq",
    "航班日期": "travel_date",
    "备注": "description",
    # English aliases
    "tour_code": "tour_code",
    "group_code": "tour_code",
    "code": "tour_code",
    "program_code": "program_code",
    "travel_date": "travel_date",
    "date": "travel_date",
    "pax": "pax",
    "size": "pax",
    "passengers": "pax",
    "amount": "amount",
    "price": "amount",
    "total": "amount",
    "fare": "unit_price",
    "unit_price": "unit_price",
    "description": "description",
    "remark": "description",
    "type": "charge_type",
    "currency": "currency",
}


def _dataframe_to_records(df: pd.DataFrame, file_type: str) -> dict:
    """Convert a DataFrame to standardized expense records."""
    field_mapping = {}

    # Map columns
    for col in df.columns:
        col_stripped = str(col).strip()
        col_lower = col_stripped.lower()
        if col_stripped in COLUMN_MAP:
            mapped = COLUMN_MAP[col_stripped]
            df.rename(columns={col: mapped}, inplace=True)
            field_mapping[col_stripped] = mapped
        elif col_lower in COLUMN_MAP:
            mapped = COLUMN_MAP[col_lower]
            df.rename(columns={col: mapped}, inplace=True)
            field_mapping[col_stripped] = mapped

    # Drop internal-only columns
    if "_seq" in df.columns:
        df.drop(columns=["_seq"], inplace=True)

    # Clean data
    if "tour_code" in df.columns:
        df["tour_code"] = df["tour_code"].astype(str).str.strip()
    if "amount" in df.columns:
        df["amount"] = pd.to_numeric(df["amount"], errors="coerce")
    if "unit_price" in df.columns:
        df["unit_price"] = pd.to_numeric(df["unit_price"], errors="coerce")
    if "pax" in df.columns:
        df["pax"] = pd.to_numeric(df["pax"], errors="coerce").fillna(0).astype(int)

    # Calculate amount from unit_price * pax if amount column is missing or NaN
    if "unit_price" in df.columns and "pax" in df.columns:
        if "amount" not in df.columns:
            df["amount"] = df["unit_price"] * df["pax"]
        else:
            mask = df["amount"].isna() & df["unit_price"].notna() & df["pax"].notna()
            df.loc[mask, "amount"] = df.loc[mask, "unit_price"] * df.loc[mask, "pax"]

    # Infer charge_type from Chinese description keywords
    _DESC_TO_CHARGE = {
        "机票": "flight", "票款": "flight", "航班": "flight",
        "团费": "land_tour", "地接": "land_tour",
        "签证": "visa", "保险": "insurance",
        "酒店": "accommodation", "住宿": "accommodation",
        "餐费": "meal", "车费": "transport",
        "门票": "entrance_fee", "导游": "guide_tip", "小费": "guide_tip",
        "服务费": "service_fee", "单房差": "single_supplement",
    }
    if "description" in df.columns and ("charge_type" not in df.columns or df["charge_type"].isna().all()):
        def _infer_charge(desc):
            if not isinstance(desc, str):
                return "other"
            for keyword, ct in _DESC_TO_CHARGE.items():
                if keyword in desc:
                    return ct
            return "other"
        df["charge_type"] = df["description"].apply(_infer_charge)

    # Set defaults
    for col, default in [("currency", "THB"), ("charge_type", "other"), ("exchange_rate", 1.0)]:
        if col not in df.columns:
            df[col] = default

    # Drop continuation rows (NaN tour_code) -- e.g., multi-leg flights
    if "tour_code" in df.columns:
        df = df[~df["tour_code"].isin(["nan", "None", ""])].copy()

    # Validate and collect records
    records = []
    errors = []
    for idx, row in df.iterrows():
        record = row.to_dict()

        row_errors = []
        tc = str(record.get("tour_code", "")).strip()
        if not tc or tc in ("nan", "None", ""):
            row_errors.append("Missing tour_code")
        if pd.isna(record.get("amount")) or record.get("amount", 0) <= 0:
            row_errors.append(f"Invalid amount: {record.get('amount')}")

        if row_errors:
            errors.append(f"Row {idx + 1}: {'; '.join(row_errors)}")
        else:
            clean_record = {}
            for k, v in record.items():
                if pd.isna(v):
                    clean_record[k] = None
                else:
                    clean_record[k] = v
            records.append(clean_record)

    return {
        "status": "success",
        "file_type": file_type,
        "columns_found": list(df.columns),
        "field_mapping": field_mapping,
        "total_rows": len(df),
        "valid_records": len(records),
        "invalid_records": len(errors),
        "records": records,
        "errors": errors,
    }


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------
def _parse_pdf(file_path: str) -> dict:
    """Extract text from PDF and use LLM to identify expense fields."""
    try:
        import pdfplumber
    except ImportError:
        return {"status": "error", "errors": ["pdfplumber not installed"], "records": []}

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            # Also try extracting tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    text_parts.append(" | ".join(str(cell or "") for cell in row))

    raw_text = "\n".join(text_parts)
    if not raw_text.strip():
        return {"status": "error", "errors": ["No text extracted from PDF"], "records": [], "raw_text": ""}

    return _extract_records_with_llm(raw_text, "pdf")


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------
def _parse_docx(file_path: str) -> dict:
    """Extract text from DOCX and use LLM to identify expense fields."""
    try:
        from docx import Document
    except ImportError:
        return {"status": "error", "errors": ["python-docx not installed"], "records": []}

    doc = Document(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    def _extract_table(table, depth=0):
        """Recursively extract text from tables, including nested tables."""
        prefix = "  " * depth
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text_parts.append(prefix + " | ".join(cells))
            # Also check for nested tables inside each cell
            for cell in row.cells:
                for nested_table in cell.tables:
                    _extract_table(nested_table, depth + 1)

    for table in doc.tables:
        _extract_table(table)

    # Also extract text from headers/footers
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        except Exception:
            pass

    raw_text = "\n".join(text_parts)
    logger.info("DOCX extracted %d text parts, %d chars total", len(text_parts), len(raw_text))
    if len(raw_text) < 200:
        logger.info("DOCX full text: %s", raw_text)
    else:
        logger.info("DOCX first 500 chars: %s", raw_text[:500])

    if not raw_text.strip():
        return {"status": "error", "errors": ["No text extracted from DOCX"], "records": [], "raw_text": ""}

    return _extract_records_with_llm(raw_text, "docx")


# ---------------------------------------------------------------------------
# Plain text parsing
# ---------------------------------------------------------------------------
def _parse_text(file_path: str) -> dict:
    """Parse plain text file using LLM for field extraction."""
    with open(file_path, "r", encoding="utf-8") as f:
        raw_text = f.read()

    if not raw_text.strip():
        return {"status": "error", "errors": ["Empty file"], "records": [], "raw_text": ""}

    return _extract_records_with_llm(raw_text, "text")


# ---------------------------------------------------------------------------
# LLM-based extraction for unstructured documents
# ---------------------------------------------------------------------------
def _extract_records_with_llm(raw_text: str, file_type: str) -> dict:
    """Use OpenAI to extract structured expense records from raw text."""
    client = OpenAI(api_key=Config.OPENAI_API_KEY, timeout=90.0)

    prompt = f"""Extract expense/charge records from the following document text.

## SUPPLIER NAME (CRITICAL)
Identify the **supplier name** — the company or person to PAY.
- Do NOT return generic contract roles like "甲方" (Party A) or "乙方" (Party B).
- For Chinese documents: find the real company name (e.g., "昆明XX旅行社").
- For Thai invoices: the supplier is the company whose BANK ACCOUNT appears at the bottom
  (e.g., "บริษัท เบสท์อินเตอร์เนชั่นแนลทราเวลแอนด์เอเจนซี่ จำกัด").
  The company in the HEADER may be the *issuer/wholesaler*, not the pay-to party.
  When in doubt, use the company with the bank account as supplier_name.
- If you cannot find a specific name, return an empty string "".

## CURRENCY DETECTION (CRITICAL)
Detect the ACTUAL currency from the document. Do NOT default to THB.
Look for these clues:
- "RMB", "人民币", "元", "¥", "CNY" -> currency = "CNY"
- "THB", "บาท", "฿", "泰铢" -> currency = "THB"
- "USD", "$", "美元", "美金" -> currency = "USD"
- "EUR", "€", "欧元" -> currency = "EUR"
- "JPY", "円", "日元" -> currency = "JPY"
- "KRW", "₩", "韩元" -> currency = "KRW"
- If the amount summary shows "(RMB)" or "元" or mentions Chinese bank accounts only -> "CNY"
- If the supplier is a Chinese company with Chinese bank accounts -> likely "CNY"
- If the document is clearly in Thai with Thai Baht amounts -> "THB"

## EXPENSE LINE ITEMS
Extract EACH expense line item as a SEPARATE record.

### Chinese expense types:
- 团费 / tour fare -> charge_type = "land_tour", expense_label = "Tour Fare"
- 单房差 / single room supplement -> charge_type = "single_supplement", expense_label = "Single Room Supplement"
- 服务费 / service fee -> charge_type = "service_fee", expense_label = "Service Fee"
- 导游费 / guide fee / 小费 / tips / 导服费 -> charge_type = "guide_tip", expense_label = "Guide Fee / Tips"
- 机票 / 票价 / flight / airfare / 机票定金 / 机票押金 -> charge_type = "flight", expense_label = "Airline Ticket"
- 签证费 / visa fee -> charge_type = "visa", expense_label = "Visa Fee"
- 酒店 / hotel / 住宿 / accommodation -> charge_type = "accommodation", expense_label = "Accommodation"
- 餐费 / meals -> charge_type = "meal", expense_label = "Meals"
- 车费 / transport / 包车 -> charge_type = "transport", expense_label = "Transport"
- 保险 / insurance -> charge_type = "insurance", expense_label = "Insurance"
- 门票 / entrance / tickets -> charge_type = "entrance_fee", expense_label = "Entrance Fee"
- 机票定金 / deposit -> use original type + "(Deposit)" suffix in expense_label

### Thai expense types:
- ค่าทัวร์ / ค่าทัวร์ผู้ใหญ่ / ค่าทัวร์เด็ก -> charge_type = "land_tour", expense_label = "Tour Fare (Adult)" or "Tour Fare (Child)"
- ค่าตั๋วเครื่องบิน / ค่ามัดจำตั๋วเครื่องบิน -> charge_type = "flight", expense_label = "Airline Ticket" or "Airline Ticket (Deposit)"
- ค่าบัตร / ค่าบัตรดีสนีย์แลนด์ / ค่าบัตรเข้าชม -> charge_type = "entrance_fee", expense_label = "Disneyland Ticket (Adult)" / use the actual attraction name
- ค่าคอมมิชชั่น / หักค่าคอมมิชชั่น -> charge_type = "commission", expense_label = "Commission"
  IMPORTANT: Commission lines are DEDUCTIONS with NEGATIVE amounts. Use the negative sign.
- ค่าบริการ -> charge_type = "service_fee", expense_label = "Service Fee"
- ค่าวีซ่า -> charge_type = "visa", expense_label = "Visa Fee"
- ค่าโรงแรม / ค่าที่พัก -> charge_type = "accommodation", expense_label = "Accommodation"
- ค่าอาหาร -> charge_type = "meal", expense_label = "Meals"
- ค่ารถ / ค่ารถบัส -> charge_type = "transport", expense_label = "Transport"
- ค่าประกัน -> charge_type = "insurance", expense_label = "Insurance"
- ค่ามัคคุเทศก์ / ค่าทิป -> charge_type = "guide_tip", expense_label = "Guide Fee / Tips"
- ค่าห้องพักเดี่ยว -> charge_type = "single_supplement", expense_label = "Single Room Supplement"

If unclear, use charge_type = "other", expense_label = translate the description to English.

## COMMISSION / DEDUCTION HANDLING
Commission lines (ค่าคอมมิชชั่น, 佣金) are DEDUCTIONS — they have NEGATIVE amounts.
- Include them as records with charge_type = "commission"
- The amount MUST be negative (e.g., -2500, not 2500)
- If the document shows "หักค่าคอมมิชชั่น 500 x 5 = -2,500" -> amount = -2500, unit_price = -500, pax = 5

## DATE EXTRACTION
- Extract BOTH the start and end dates of travel if available.
- travel_date_start: First day of travel in dd/mm/yyyy format (e.g., "04/03/2026")
- travel_date_end: Last day of travel in dd/mm/yyyy format (e.g., "09/03/2026")
- travel_date: The original date text as-is from the document (for reference)

COMMON DATE FORMATS:
- "0304-0309" -> mmdd-mmdd -> March 4 to March 9
- "3月4日-9日" -> March 4-9 (Chinese)
- "0622-0626" -> June 22 to June 26
- The YEAR comes from the tour code. If the tour code ends with 260304 (yymmdd),
  the year is 20yy = 2026.

### Thai dates with Buddhist Era (พ.ศ.)
Thai years use the Buddhist Era (พ.ศ.), which is 543 years ahead of CE.
- พ.ศ. 2569 = CE 2026  (2569 - 543 = 2026)
- พ.ศ. 2568 = CE 2025
- Month abbreviations: ม.ค.=Jan, ก.พ.=Feb, มี.ค.=Mar, เม.ย.=Apr, พ.ค.=May,
  มิ.ย.=Jun, ก.ค.=Jul, ส.ค.=Aug, ก.ย.=Sep, ต.ค.=Oct, พ.ย.=Nov, ธ.ค.=Dec
- Example: "13 มี.ค. 2569 - 18 มี.ค. 2569" -> travel_date_start = "13/03/2026", travel_date_end = "18/03/2026"
- Example: "วันเดินทาง : 13 มี.ค. 2569" -> start date = 13/03/2026

## TOUR CODE / GROUP CODE
Look for these keywords:
- "GROUP CODE", "Group Code", "รหัสกรุ๊ป", "Code group"
- Chinese: "团号", "组号", "合同号"
- The value after the keyword is the tour_code (e.g., "GROUP CODE : BTNRTXJ260313W02")

## GROUPING (CRITICAL)
All line items from the SAME document that belong to the same tour/trip MUST share
the EXACT same tour_code. Do NOT create different tour_code values for items that
belong together (e.g., airline ticket deposit + tour fare + commission for the same trip).
If the document only mentions ONE group/tour code, ALL records must use that SAME code.

EXCEPTION — MULTI-CODE INVOICES:
Some invoices list MULTIPLE separate tour codes in a table (e.g., 4 different flight
bookings each with their own code like GO1TAO5NTAOQW260118, GO1TAO5NTAOQW260121, etc.).
In this case, each row with a DIFFERENT tour code IS a separate group. Preserve the
distinct tour_code for each row. Do NOT merge them under one code.

## COST CALCULATION STRUCTURE (CRITICAL)
Preserve HOW the amount is calculated. Many invoices show multi-factor math:
- Simple: unit_price x pax = amount (e.g., 2,900 x 17 = 49,300)
- With days: rate x pax x days = amount (e.g., tips 30 x 20 x 7 = 4,200)
- With quantity: unit_price x pax x quantity = amount (e.g., 20 x 16 x 6 = 1,920)
- Deductions: negative amounts (e.g., -800 x 1 = -800)

For EACH line item, extract:
- tour_code: Tour or group code. ALL items for the same trip MUST have the EXACT same tour_code string.
- program_code: Program code if available (e.g., "BT-NRT_W02_XJ")
- travel_date: Original date text as-is
- travel_date_start: Start date in dd/mm/yyyy or null
- travel_date_end: End date in dd/mm/yyyy or null
- pax: Number of passengers / units for THIS line item
- unit_price: Price per person/unit. MUST extract this.
- quantity: Extra multiplier beyond pax (e.g., number of days, nights, meals). null if not applicable.
- amount: Total for this line (unit_price * pax * quantity, or unit_price * pax). Use NEGATIVE for commissions/deductions.
- currency: The detected currency code
- description: Original description text (Chinese, Thai, or English as-is)
- expense_label: Human-readable ENGLISH label for the expense type. Translate from Chinese/Thai if needed.
- charge_type: Category code from the lists above
- calculation_note: Human-readable string showing the math as it appears in the invoice.
  Examples: "2,380 x 20 pax = 47,600", "30 x 20 pax x 7 days = 4,200",
  "650 x 2 rooms = 1,300", "-800 x 1 = -800". ALWAYS include this.

Return ONLY valid JSON in this format:
{{
    "supplier_name": "the company/supplier name (pay-to party)",
    "detected_currency": "CNY or THB or USD etc.",
    "currency_evidence": "brief explanation of how you detected the currency",
    "records": [
        {{
            "tour_code": "string",
            "program_code": "string or null",
            "travel_date": "string or null",
            "travel_date_start": "dd/mm/yyyy or null",
            "travel_date_end": "dd/mm/yyyy or null",
            "pax": number or null,
            "unit_price": number or null,
            "quantity": number or null,
            "amount": number,
            "currency": "the detected currency",
            "description": "string (original text)",
            "expense_label": "string (English label, e.g. Airline Ticket)",
            "charge_type": "string from categories above",
            "calculation_note": "string showing the math (e.g. '2,380 x 20 pax = 47,600')"
        }}
    ],
    "notes": "any observations about the data"
}}

Document text:
---
{raw_text[:6000]}
---"""

    try:
        response = client.chat.completions.create(
            model=Config.OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You extract structured data from documents. Return only valid JSON.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.1,
            max_tokens=4000,
            response_format={"type": "json_object"},
        )

        raw_content = response.choices[0].message.content
        result = json.loads(raw_content)
        records = result.get("records", [])
        supplier_name = result.get("supplier_name", "")
        detected_currency = result.get("detected_currency", "")
        currency_evidence = result.get("currency_evidence", "")

        logger.info("LLM returned %d records, supplier='%s'", len(records), supplier_name)
        if detected_currency:
            logger.info("Currency detected: %s (evidence: %s)", detected_currency, currency_evidence)
        if not records:
            logger.warning("LLM returned 0 records! Raw response (first 500): %s", raw_content[:500])

        for rec in records:
            rec["supplier_name"] = supplier_name
            if detected_currency and rec.get("currency") in ("THB", None, ""):
                rec["currency"] = detected_currency

        return {
            "status": "success",
            "file_type": file_type,
            "extraction_method": "llm",
            "supplier_name": supplier_name,
            "detected_currency": detected_currency,
            "currency_evidence": currency_evidence,
            "total_rows": len(records),
            "valid_records": len(records),
            "invalid_records": 0,
            "records": records,
            "raw_text": raw_text[:2000],
            "notes": result.get("notes", ""),
            "errors": [],
        }

    except Exception as e:
        logger.error(f"LLM extraction failed: {e}", exc_info=True)
        return {
            "status": "error",
            "file_type": file_type,
            "raw_text": raw_text[:2000],
            "errors": [f"AI extraction failed: {str(e)}"],
            "records": [],
        }
